import pandas as pd
import docx
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime
import os

time = datetime.datetime.now()
timeformat = time.strftime("%Y-%m-%d_%H-%M-%S")

# Load the Excel file
excel_file = 'data.xlsx'
df = pd.read_excel(excel_file)

# Load the Word template
template_file = 'template.docx'

# Define the mail merge fields including the new ones
fields = [
    'Organization', 'Name', 'Address', 'City', 'State', 'Zip', 'Phone', 'Email', 'Date', 'Assessment', 'Dates', 
    'Assessment_location', 'Scope', 'Standards', 'Findings', 'Evidence', 'Scope_Specifics', 'Risk', 'Impact', 
    'Risk_Result', 'Recommendations'
]

# Function to format date fields
def format_date(date_value):
    if isinstance(date_value, pd.Timestamp):
        return date_value.strftime('%Y-%m-%d')  # Format as YYYY-MM-DD
    return date_value

# Function to create an invisible table with headers and data
def add_invisible_table(doc, findings_data, evidence_data, risk_data, impact_data, risk_result_data, recommendations_data, insert_after_paragraph):
    findings = findings_data.split(',')
    evidences = evidence_data.split(',')
    risks = risk_data.split(',')
    impacts = impact_data.split(',')
    risk_results = risk_result_data.split(',')
    recommendations = recommendations_data.split(',')

    # Make sure all lists have the same length by padding with empty strings if needed
    max_length = max(len(findings), len(evidences), len(risks), len(impacts), len(risk_results), len(recommendations))
    findings += [''] * (max_length - len(findings))
    evidences += [''] * (max_length - len(evidences))
    risks += [''] * (max_length - len(risks))
    impacts += [''] * (max_length - len(impacts))
    risk_results += [''] * (max_length - len(risk_results))
    recommendations += [''] * (max_length - len(recommendations))

    # Create a table
    table = doc.add_table(rows=1, cols=6)  # Add a table with one header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Finding'
    hdr_cells[1].text = 'Evidence'
    hdr_cells[2].text = 'Risk'
    hdr_cells[3].text = 'Impact'
    hdr_cells[4].text = 'Risk Result'
    hdr_cells[5].text = 'Recommendations'

    # Bold the headers
    for cell in hdr_cells:
        cell.text = cell.text.strip()  # Clean up any leading/trailing whitespace
        cell.paragraphs[0].runs[0].bold = True  # Bold the header text

    # Add data to the table
    for finding, evidence, risk, impact, risk_result, recommendation in zip(findings, evidences, risks, impacts, risk_results, recommendations):
        row_cells = table.add_row().cells
        row_cells[0].text = finding.strip()
        
        # Handle Evidence - check if it's an image file
        if os.path.exists(evidence.strip()):
            paragraph = row_cells[1].paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(evidence.strip(), width=Inches(1.5))
        else:
            row_cells[1].text = evidence.strip()

        row_cells[2].text = risk.strip()
        row_cells[3].text = impact.strip()
        row_cells[4].text = risk_result.strip()
        row_cells[5].text = recommendation.strip()

    # Set border to 1 pixel
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_pr.append(parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4" w_space="0"/><w:left w:val="single" w:sz="4" w_space="0"/><w:bottom w:val="single" w:sz="4" w_space="0"/><w:right w:val="single" w:sz="4" w_space="0"/><w:insideH w:val="single" w:sz="4" w_space="0"/><w:insideV w:val="single" w:sz="4" w_space="0"/></w:tblBorders>' % nsdecls('w')))

    # Move the table right after the paragraph that contained the placeholder
    doc.element.body.insert(doc.element.body.index(insert_after_paragraph._element) + 1, tbl)

# Function to handle Scope_Specifics
def add_scope_specifics_table(doc, scope_specifics_data, insert_after_paragraph):
    scope_specifics = scope_specifics_data.split(',')

    # Create a table for Scope Specifics
    table = doc.add_table(rows=1, cols=1)  # Add a table with one column
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Scope Specifics'

    # Bold the header
    hdr_cells[0].paragraphs[0].runs[0].bold = True  # Bold the header text

    # Add each Scope Specific as a new row
    for scope_specific in scope_specifics:
        row_cells = table.add_row().cells
        row_cells[0].text = scope_specific.strip()

    # Set border to 1 pixel
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_pr.append(parse_xml(r'<w:tblBorders %s><w:top w:val="single" w:sz="4" w_space="0"/><w:left w:val="single" w:sz="4" w_space="0"/><w:bottom w:val="single" w:sz="4" w_space="0"/><w:right w:val="single" w:sz="4" w_space="0"/><w:insideH w:val="single" w:sz="4" w_space="0"/><w:insideV w:val="single" w:sz="4" w_space="0"/></w:tblBorders>' % nsdecls('w')))

    # Move the table right after the paragraph that contained the placeholder
    doc.element.body.insert(doc.element.body.index(insert_after_paragraph._element) + 1, tbl)


# Iterate over the Excel data and perform the mail merge
for index, row in df.iterrows():
    doc = docx.Document(template_file)

    for paragraph in doc.paragraphs:
        for field in fields:
            placeholder = '{{ ' + field + ' }}'
            if placeholder in paragraph.text:
                # Format the Date and Dates fields
                if field in ['Date', 'Dates']:
                    formatted_date = format_date(row[field])
                    paragraph.text = paragraph.text.replace(placeholder, str(formatted_date))

                # Handle Findings, Evidence, Risk, Impact, Risk Result, and Recommendations together in a table
                elif field == 'Findings' and isinstance(row['Findings'], str) and isinstance(row['Evidence'], str) and isinstance(row['Risk'], str) and isinstance(row['Impact'], str) and isinstance(row['Risk_Result'], str) and isinstance(row['Recommendations'], str):
                    paragraph.text = paragraph.text.replace(placeholder, "")
                    add_invisible_table(doc, row['Findings'], row['Evidence'], row['Risk'], row['Impact'], row['Risk_Result'], row['Recommendations'], paragraph)
                
                # Handle Scope_Specifics field as a new row for each entry
                elif field == 'Scope_Specifics' and isinstance(row['Scope_Specifics'], str):
                    paragraph.text = paragraph.text.replace(placeholder, "")
                    add_scope_specifics_table(doc, row['Scope_Specifics'], paragraph)
                    paragraph.add_run().add_break()
                    
                # Handle other fields
                else:
                    paragraph.text = paragraph.text.replace(placeholder, str(row[field]))

    # Save the result
    try:
        output_file = f'output_{index}_{timeformat}.docx'
        doc.save(output_file)
        print(f"File saved successfully: {output_file}")
    except OSError as e:
        print(f"Error saving file: {e}")
